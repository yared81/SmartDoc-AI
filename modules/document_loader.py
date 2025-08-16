import os
import io
from typing import List
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader, UnstructuredWordDocumentLoader, UnstructuredExcelLoader
import pandas as pd
from docx import Document as DocxDocument

def load_pdf(file_content: bytes, filename: str) -> List[Document]:
    """Load PDF documents from file content."""
    try:
        # Save temporary file
        temp_path = f"temp_{filename}"
        with open(temp_path, "wb") as f:
            f.write(file_content)
        
        # Load with PyPDFLoader
        loader = PyPDFLoader(temp_path)
        docs = loader.load()
        
        # Clean up
        os.remove(temp_path)
        
        # Add metadata
        for doc in docs:
            doc.metadata["source"] = filename
            doc.metadata["type"] = "pdf"
        
        return docs
    except Exception as e:
        raise RuntimeError(f"Failed to load PDF {filename}: {e}")

def load_docx(file_content: bytes, filename: str) -> List[Document]:
    """Load DOCX documents from file content."""
    try:
        # Load with python-docx
        doc = DocxDocument(io.BytesIO(file_content))
        
        # Extract text from paragraphs
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        
        # Create document
        doc_obj = Document(
            page_content=text,
            metadata={"source": filename, "type": "docx"}
        )
        
        return [doc_obj]
    except Exception as e:
        raise RuntimeError(f"Failed to load DOCX {filename}: {e}")

def load_excel(file_content: bytes, filename: str) -> List[Document]:
    """Load Excel documents from file content."""
    try:
        # Load with pandas
        df = pd.read_excel(io.BytesIO(file_content), sheet_name=None)
        
        docs = []
        for sheet_name, sheet_df in df.items():
            # Convert DataFrame to text
            text = f"Sheet: {sheet_name}\n\n"
            text += sheet_df.to_string(index=False)
            
            doc = Document(
                page_content=text,
                metadata={"source": filename, "type": "excel", "sheet": sheet_name}
            )
            docs.append(doc)
        
        return docs
    except Exception as e:
        raise RuntimeError(f"Failed to load Excel {filename}: {e}")

def load_text(file_content: bytes, filename: str) -> List[Document]:
    """Load text documents from file content."""
    try:
        text = file_content.decode('utf-8')
        
        doc = Document(
            page_content=text,
            metadata={"source": filename, "type": "text"}
        )
        
        return [doc]
    except Exception as e:
        raise RuntimeError(f"Failed to load text {filename}: {e}")

def load_documents(file_content: bytes, filename: str) -> List[Document]:
    """Load documents based on file extension."""
    file_extension = filename.lower().split('.')[-1]
    
    if file_extension == 'pdf':
        return load_pdf(file_content, filename)
    elif file_extension in ['doc', 'docx']:
        return load_docx(file_content, filename)
    elif file_extension in ['xls', 'xlsx']:
        return load_excel(file_content, filename)
    elif file_extension == 'txt':
        return load_text(file_content, filename)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

def load_from_text(text_content: str, source_name: str = "Manual Input") -> List[Document]:
    """Load documents from manually entered text."""
    doc = Document(
        page_content=text_content,
        metadata={"source": source_name, "type": "manual"}
    )
    return [doc]
